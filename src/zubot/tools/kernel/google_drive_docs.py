"""DOCX generation and Google Drive upload helpers."""

from __future__ import annotations

import json
import mimetypes
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote, urlencode
from urllib.request import Request, urlopen

from src.zubot.core.config_loader import load_config
from src.zubot.tools.kernel.google_auth import get_google_access_token

DEFAULT_TIMEOUT_SEC = 15
DEFAULT_UPLOAD_PATH = "Job Applications/Cover Letters"
DEFAULT_OUTPUT_DIR = "outputs/cover_letters"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
FOLDER_MIME_TYPE = "application/vnd.google-apps.folder"


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[4]


def _resolve_repo_relative_path(raw_path: str) -> Path:
    candidate = Path(raw_path.strip())
    if candidate.is_absolute():
        raise ValueError("Absolute paths are not allowed; use repository-relative paths.")
    normalized_parts = [part for part in candidate.parts if part not in (".", "")]
    if any(part == ".." for part in normalized_parts):
        raise ValueError("Path traversal is not allowed.")
    return (_repo_root() / candidate).resolve()


def _ensure_docx_filename(filename: str) -> str:
    base = filename.strip()
    if not base:
        raise ValueError("filename must be non-empty.")
    if "/" in base or "\\" in base:
        raise ValueError("filename must not contain path separators.")
    if base.lower().endswith(".docx"):
        return base
    return f"{base}.docx"


def _google_drive_settings() -> dict[str, Any]:
    try:
        payload = load_config()
    except (FileNotFoundError, ValueError):
        payload = {}

    config: dict[str, Any] = {}
    profiles = payload.get("tool_profiles")
    if isinstance(profiles, dict):
        user_specific = profiles.get("user_specific")
        if isinstance(user_specific, dict):
            block = user_specific.get("google_drive")
            if isinstance(block, dict):
                config = block

    spreadsheet_id = config.get("job_application_spreadsheet_id")
    default_upload_path = config.get("default_upload_path")
    cover_letters_folder_id = config.get("cover_letters_folder_id")
    timeout_sec = config.get("timeout_sec", DEFAULT_TIMEOUT_SEC)

    return {
        "job_application_spreadsheet_id": spreadsheet_id if isinstance(spreadsheet_id, str) else None,
        "default_upload_path": default_upload_path if isinstance(default_upload_path, str) and default_upload_path.strip() else DEFAULT_UPLOAD_PATH,
        "cover_letters_folder_id": cover_letters_folder_id.strip()
        if isinstance(cover_letters_folder_id, str) and cover_letters_folder_id.strip()
        else None,
        "timeout_sec": int(timeout_sec),
    }


def _authorized_headers(access_token: str) -> dict[str, str]:
    return {
        "Authorization": f"Bearer {access_token}",
        "Accept": "application/json",
    }


def _fetch_json(url: str, headers: dict[str, str], timeout_sec: int) -> dict[str, Any]:
    req = Request(url, headers=headers, method="GET")
    with urlopen(req, timeout=timeout_sec) as response:
        body = response.read().decode("utf-8")
    payload = json.loads(body)
    if not isinstance(payload, dict):
        raise ValueError("Google Drive response must be a JSON object.")
    return payload


def _post_json(url: str, headers: dict[str, str], payload: dict[str, Any], timeout_sec: int) -> dict[str, Any]:
    request_headers = dict(headers)
    request_headers["Content-Type"] = "application/json"
    req = Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers=request_headers,
        method="POST",
    )
    with urlopen(req, timeout=timeout_sec) as response:
        body = response.read().decode("utf-8")
    data = json.loads(body)
    if not isinstance(data, dict):
        raise ValueError("Google Drive response must be a JSON object.")
    return data


def _upload_multipart(
    *,
    access_token: str,
    metadata: dict[str, Any],
    content: bytes,
    content_type: str,
    timeout_sec: int,
) -> dict[str, Any]:
    boundary = "zubotBoundary7MA4YWxkTrZu0gW"
    parts = [
        f"--{boundary}\r\n".encode("utf-8"),
        b"Content-Type: application/json; charset=UTF-8\r\n\r\n",
        json.dumps(metadata).encode("utf-8"),
        b"\r\n",
        f"--{boundary}\r\n".encode("utf-8"),
        f"Content-Type: {content_type}\r\n\r\n".encode("utf-8"),
        content,
        b"\r\n",
        f"--{boundary}--\r\n".encode("utf-8"),
    ]
    body = b"".join(parts)

    url = "https://www.googleapis.com/upload/drive/v3/files?uploadType=multipart&fields=id,name,webViewLink,mimeType,parents"
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Accept": "application/json",
        "Content-Type": f"multipart/related; boundary={boundary}",
    }
    req = Request(url, data=body, headers=headers, method="POST")
    with urlopen(req, timeout=timeout_sec) as response:
        raw = response.read().decode("utf-8")
    payload = json.loads(raw)
    if not isinstance(payload, dict):
        raise ValueError("Google Drive upload response must be a JSON object.")
    return payload


def _build_drive_list_url(query: str, fields: str) -> str:
    params = urlencode({"q": query, "fields": fields, "pageSize": 50, "spaces": "drive"})
    return f"https://www.googleapis.com/drive/v3/files?{params}"


def _escape_query_value(value: str) -> str:
    return value.replace("'", "\\'")


def _find_child_folder_id(
    *,
    access_token: str,
    parent_id: str,
    folder_name: str,
    timeout_sec: int,
) -> str | None:
    escaped_name = _escape_query_value(folder_name)
    escaped_parent = _escape_query_value(parent_id)
    query = (
        f"name = '{escaped_name}' and mimeType = '{FOLDER_MIME_TYPE}' "
        f"and '{escaped_parent}' in parents and trashed = false"
    )
    url = _build_drive_list_url(query, "files(id,name)")
    payload = _fetch_json(url, _authorized_headers(access_token), timeout_sec)
    files = payload.get("files")
    if not isinstance(files, list) or not files:
        return None
    first = files[0]
    if isinstance(first, dict) and isinstance(first.get("id"), str):
        return first["id"]
    return None


def _create_folder(
    *,
    access_token: str,
    parent_id: str,
    folder_name: str,
    timeout_sec: int,
) -> str:
    payload = {
        "name": folder_name,
        "mimeType": FOLDER_MIME_TYPE,
        "parents": [parent_id],
    }
    url = "https://www.googleapis.com/drive/v3/files?fields=id,name,mimeType,parents"
    created = _post_json(url, _authorized_headers(access_token), payload, timeout_sec)
    folder_id = created.get("id")
    if not isinstance(folder_id, str) or not folder_id:
        raise ValueError("Folder creation response missing id.")
    return folder_id


def _resolve_or_create_folder_path(*, access_token: str, path: str, timeout_sec: int) -> str:
    parts = [segment.strip() for segment in path.split("/") if segment.strip()]
    if not parts:
        raise ValueError("destination_path must include at least one folder segment.")

    parent_id = "root"
    for segment in parts:
        found = _find_child_folder_id(
            access_token=access_token,
            parent_id=parent_id,
            folder_name=segment,
            timeout_sec=timeout_sec,
        )
        if found:
            parent_id = found
            continue
        parent_id = _create_folder(
            access_token=access_token,
            parent_id=parent_id,
            folder_name=segment,
            timeout_sec=timeout_sec,
        )
    return parent_id


def _validate_folder_id(*, access_token: str, folder_id: str, timeout_sec: int) -> str:
    encoded_folder_id = quote(folder_id, safe="")
    url = f"https://www.googleapis.com/drive/v3/files/{encoded_folder_id}?fields=id,mimeType,trashed"
    payload = _fetch_json(url, _authorized_headers(access_token), timeout_sec)
    mime_type = payload.get("mimeType")
    if mime_type != FOLDER_MIME_TYPE:
        raise ValueError("Configured cover_letters_folder_id is not a folder.")
    if payload.get("trashed") is True:
        raise ValueError("Configured cover_letters_folder_id points to a trashed folder.")
    return folder_id


def _resolve_destination_folder_id(
    *,
    access_token: str,
    settings: dict[str, Any],
    destination_path: str,
    timeout_sec: int,
) -> str:
    configured_folder_id = settings.get("cover_letters_folder_id")
    if configured_folder_id and destination_path == settings["default_upload_path"]:
        try:
            return _validate_folder_id(
                access_token=access_token,
                folder_id=str(configured_folder_id),
                timeout_sec=timeout_sec,
            )
        except Exception:
            # Configured folder IDs can go stale (moved/deleted/permissions). Fall
            # back to path-based resolution so uploads still succeed.
            return _resolve_or_create_folder_path(access_token=access_token, path=destination_path, timeout_sec=timeout_sec)
    return _resolve_or_create_folder_path(access_token=access_token, path=destination_path, timeout_sec=timeout_sec)


def _file_exists_in_folder(*, access_token: str, parent_id: str, filename: str, timeout_sec: int) -> bool:
    escaped_name = _escape_query_value(filename)
    escaped_parent = _escape_query_value(parent_id)
    query = f"name = '{escaped_name}' and '{escaped_parent}' in parents and trashed = false"
    url = _build_drive_list_url(query, "files(id,name)")
    payload = _fetch_json(url, _authorized_headers(access_token), timeout_sec)
    files = payload.get("files")
    return isinstance(files, list) and len(files) > 0


def _with_timestamp_suffix(filename: str) -> str:
    stem = Path(filename).stem
    ext = Path(filename).suffix or ".docx"
    stamp = datetime.now(tz=UTC).strftime("%Y%m%d-%H%M%S")
    return f"{stem}-{stamp}{ext}"


def _error(source: str, message: str) -> dict[str, Any]:
    return {"ok": False, "source": source, "error": message}


def create_local_docx(
    *,
    filename: str,
    title: str | None = None,
    paragraphs: list[str],
    output_dir: str = DEFAULT_OUTPUT_DIR,
) -> dict[str, Any]:
    source = "google_docx_local"

    if not isinstance(paragraphs, list) or not paragraphs:
        return _error("google_docx_local_error", "paragraphs must be a non-empty list of strings.")
    if any(not isinstance(item, str) or not item.strip() for item in paragraphs):
        return _error("google_docx_local_error", "paragraphs must contain non-empty strings.")

    try:
        safe_name = _ensure_docx_filename(filename)
        output_dir_path = _resolve_repo_relative_path(output_dir)
    except ValueError as exc:
        return _error("google_docx_local_error", str(exc))

    output_dir_path.mkdir(parents=True, exist_ok=True)
    local_path = output_dir_path / safe_name

    try:
        from docx import Document  # type: ignore
    except Exception:
        return _error("google_docx_local_error", "python-docx is not available in the environment.")

    try:
        doc = Document()
        if isinstance(title, str) and title.strip():
            doc.add_heading(title.strip(), level=1)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph.strip())
        doc.save(str(local_path))
    except Exception as exc:
        return _error("google_docx_local_error", f"Failed to write DOCX: {exc}")

    bytes_written = None
    try:
        bytes_written = local_path.stat().st_size
    except OSError:
        bytes_written = None

    return {
        "ok": True,
        "source": source,
        "local_path": str(local_path.relative_to(_repo_root())),
        "filename": safe_name,
        "bytes_written": bytes_written,
        "error": None,
    }


def upload_file_to_google_drive(
    *,
    local_path: str,
    destination_path: str = DEFAULT_UPLOAD_PATH,
    destination_folder_id: str | None = None,
    filename: str | None = None,
) -> dict[str, Any]:
    source = "google_drive_upload"

    if not isinstance(local_path, str) or not local_path.strip():
        return _error("google_drive_upload_error", "local_path must be non-empty.")

    try:
        resolved_local_path = _resolve_repo_relative_path(local_path)
    except ValueError as exc:
        return _error("google_drive_upload_error", str(exc))

    if not resolved_local_path.exists() or not resolved_local_path.is_file():
        return _error("google_drive_upload_error", "local_path does not exist as a file.")

    try:
        desired_name = _ensure_docx_filename(filename or resolved_local_path.name)
    except ValueError as exc:
        return _error("google_drive_upload_error", str(exc))

    settings = _google_drive_settings()
    timeout_sec = settings["timeout_sec"]
    target_path = destination_path.strip() if isinstance(destination_path, str) and destination_path.strip() else settings["default_upload_path"]
    target_folder_id = destination_folder_id.strip() if isinstance(destination_folder_id, str) and destination_folder_id.strip() else None

    token = get_google_access_token()
    if not token.get("ok"):
        return _error("google_drive_upload_error", f"Google auth failed: {token.get('error')}")

    access_token = str(token.get("access_token") or "")
    if not access_token:
        return _error("google_drive_upload_error", "Google auth returned empty access token.")

    try:
        if target_folder_id:
            folder_id = _validate_folder_id(
                access_token=access_token,
                folder_id=target_folder_id,
                timeout_sec=timeout_sec,
            )
        else:
            folder_id = _resolve_destination_folder_id(
                access_token=access_token,
                settings=settings,
                destination_path=target_path,
                timeout_sec=timeout_sec,
            )
    except Exception as exc:
        return _error("google_drive_path_resolve_error", f"Failed to resolve destination path: {exc}")

    final_name = desired_name
    try:
        if _file_exists_in_folder(
            access_token=access_token,
            parent_id=folder_id,
            filename=desired_name,
            timeout_sec=timeout_sec,
        ):
            final_name = _with_timestamp_suffix(desired_name)
    except Exception as exc:
        return _error("google_drive_upload_error", f"Failed to check filename conflict: {exc}")

    content_type = mimetypes.guess_type(final_name)[0] or DOCX_MIME_TYPE
    try:
        content = resolved_local_path.read_bytes()
    except OSError as exc:
        return _error("google_drive_upload_error", f"Failed to read local file: {exc}")

    metadata = {"name": final_name, "parents": [folder_id], "mimeType": content_type}
    try:
        payload = _upload_multipart(
            access_token=access_token,
            metadata=metadata,
            content=content,
            content_type=content_type,
            timeout_sec=timeout_sec,
        )
    except Exception as exc:
        return _error("google_drive_upload_error", f"Failed to upload file: {exc}")

    drive_file_id = payload.get("id")
    web_view_link = payload.get("webViewLink")
    if (not isinstance(web_view_link, str) or not web_view_link.strip()) and isinstance(drive_file_id, str) and drive_file_id.strip():
        web_view_link = f"https://drive.google.com/file/d/{drive_file_id.strip()}/view"

    return {
        "ok": True,
        "source": source,
        "drive_file_id": drive_file_id,
        "drive_file_name": payload.get("name") or final_name,
        "destination_folder_id": folder_id,
        "drive_folder_id": folder_id,
        "web_view_link": web_view_link,
        "error": None,
    }


def create_and_upload_docx(
    *,
    filename: str,
    title: str | None = None,
    paragraphs: list[str],
    output_dir: str = DEFAULT_OUTPUT_DIR,
    destination_path: str = DEFAULT_UPLOAD_PATH,
) -> dict[str, Any]:
    local_result = create_local_docx(filename=filename, title=title, paragraphs=paragraphs, output_dir=output_dir)
    if not local_result.get("ok"):
        return {
            "ok": False,
            "source": "google_docx_create_upload_error",
            "error": local_result.get("error"),
            "local": local_result,
            "upload": None,
        }

    upload_result = upload_file_to_google_drive(
        local_path=str(local_result.get("local_path") or ""),
        destination_path=destination_path,
        filename=str(local_result.get("filename") or ""),
    )
    if not upload_result.get("ok"):
        return {
            "ok": False,
            "source": "google_docx_create_upload_error",
            "error": upload_result.get("error"),
            "local": local_result,
            "upload": upload_result,
        }

    return {
        "ok": True,
        "source": "google_docx_create_upload",
        "local": local_result,
        "upload": upload_result,
        "error": None,
    }
