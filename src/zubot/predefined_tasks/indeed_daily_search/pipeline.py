"""Indeed daily search task pipeline."""

from __future__ import annotations

from concurrent.futures import FIRST_COMPLETED, Future, ThreadPoolExecutor, wait
from dataclasses import dataclass
from datetime import UTC, datetime
import hashlib
import json
from pathlib import Path
from queue import Queue
import re
import sqlite3
from threading import Event, Lock, Thread
from time import sleep
from typing import Any, Callable
from urllib.parse import parse_qs, urlparse
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

from src.zubot.core.config_loader import get_central_service_config, get_timezone, load_config
from src.zubot.core.llm_client import call_llm
from src.zubot.core.task_scheduler_store import resolve_scheduler_db_path
from src.zubot.tools.kernel.google_drive_docs import upload_file_to_google_drive
from src.zubot.tools.kernel.google_sheets_job_apps import append_job_app_row
from src.zubot.tools.kernel.hasdata_indeed import get_indeed_job_detail, get_indeed_jobs

DECISION_RECOMMEND_APPLY = "Recommend Apply"
DECISION_RECOMMEND_MAYBE = "Recommend Maybe"
DECISION_SKIP = "Skip"
ALLOWED_DECISIONS = {DECISION_RECOMMEND_APPLY, DECISION_RECOMMEND_MAYBE, DECISION_SKIP}

DEFAULT_SEEN_LIMIT = 200
DEFAULT_PROVIDER = "indeed"
NOT_FOUND_VALUE = "Not Found"
DEFAULT_EXTRACTION_MODEL = "medium"
DEFAULT_DECISION_MODEL = "medium"
DEFAULT_COVER_LETTER_MODEL = "high"
DEFAULT_COVER_LETTER_DESTINATION_PATH = "Job Applications/Cover Letters"
DEFAULT_FILE_MODE = "versioned"
DEFAULT_COVER_LETTER_CONTACT_EMAIL = "zubinkjha2025@gmail.com"
DEFAULT_COVER_LETTER_LINKEDIN_URL = "https://www.linkedin.com/in/zubin-jha-30752a355"
DEFAULT_COVER_LETTER_LINKEDIN_LABEL = "LinkedIn"
DEFAULT_COVER_LETTER_UPLOAD_RETRY_ATTEMPTS = 3
DEFAULT_COVER_LETTER_UPLOAD_RETRY_BACKOFF_SEC = 1.0
DEFAULT_SHEET_RETRY_ATTEMPTS = 2
DEFAULT_SHEET_RETRY_BACKOFF_SEC = 1.0
DEFAULT_PROJECT_CONTEXT_TOP_N = 3
DEFAULT_PROJECT_CONTEXT_MAX_CHARS = 2600
DEFAULT_PROCESS_WORKERS = 1
DEFAULT_DB_QUEUE_MAXSIZE = 128
DEFAULT_SHEET_QUEUE_MAXSIZE = 128
MAX_PROCESS_WORKERS = 12
SEARCH_PHASE_WEIGHT = 0.02

DEFAULT_CANDIDATE_CONTEXT_FILES = [
    "context/USER.md",
    "context/more-about-human/README.md",
    "context/more-about-human/professional_profile.md",
    "context/more-about-human/passion_profile.md",
    "context/more-about-human/writing_voice.md",
    "context/more-about-human/resume.md",
    "context/more-about-human/more-about-me.md",
    "context/more-about-human/cover_letter_brain.json",
    "context/more-about-human/projects/project_index.md",
]

_KEY_CHARS_PATTERN = re.compile(r"[^a-zA-Z0-9._-]+")
_TOKEN_PATTERN = re.compile(r"[a-z0-9]{3,}")
_WORD_PATTERN = re.compile(r"[A-Za-z0-9']+")


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[4]


def _utc_now() -> datetime:
    return datetime.now(tz=UTC)


def _iso_now() -> str:
    return _utc_now().isoformat()


def _local_today_iso() -> str:
    tz_name = "America/New_York"
    try:
        cfg = load_config()
        configured_tz = get_timezone(cfg)
        if isinstance(configured_tz, str) and configured_tz.strip():
            tz_name = configured_tz.strip()
    except Exception:
        pass
    try:
        return datetime.now(ZoneInfo(tz_name)).date().isoformat()
    except ZoneInfoNotFoundError:
        return datetime.now(ZoneInfo("America/New_York")).date().isoformat()
    except Exception:
        return _utc_now().date().isoformat()


def _coerce_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _extract_from_dict(payload: dict[str, Any], keys: list[str]) -> str:
    for key in keys:
        value = payload.get(key)
        text = _coerce_text(value)
        if text:
            return text
    return ""


def _extract_job_key(job: dict[str, Any]) -> str:
    explicit = _extract_from_dict(job, ["jobKey", "jobkey", "job_key", "key", "id"])
    if explicit:
        return explicit
    url = _extract_job_url(job)
    if url:
        try:
            parsed = urlparse(url)
            query = parse_qs(parsed.query)
            jk = query.get("jk")
            if isinstance(jk, list) and jk and _coerce_text(jk[0]):
                return _coerce_text(jk[0])
        except Exception:
            pass
    material = "|".join(
        [
            _extract_job_title(job),
            _extract_job_company(job),
            _extract_job_location(job),
            url,
        ]
    )
    if not material.strip():
        material = json.dumps(job, ensure_ascii=True, sort_keys=True)
    return hashlib.sha1(material.encode("utf-8")).hexdigest()[:20]


def _extract_job_url(job: dict[str, Any]) -> str:
    return _extract_from_dict(job, ["url", "jobUrl", "job_url", "jobLink", "link"])


def _extract_job_title(job: dict[str, Any]) -> str:
    return _extract_from_dict(job, ["title", "jobTitle", "job_title", "position"])


def _extract_job_company(job: dict[str, Any]) -> str:
    company = job.get("company")
    if isinstance(company, dict):
        nested = _extract_from_dict(company, ["name", "companyName"])
        if nested:
            return nested
    return _extract_from_dict(job, ["companyName", "company", "company_name", "employer"])


def _extract_job_location(job: dict[str, Any]) -> str:
    location = job.get("location")
    if isinstance(location, dict):
        nested = _extract_from_dict(location, ["formattedAddress", "displayName", "name"])
        if nested:
            return nested
    return _extract_from_dict(job, ["location", "jobLocation", "formattedLocation", "cityState"])


def _sanitize_file_stem(value: str) -> str:
    clean = _KEY_CHARS_PATTERN.sub("_", value.strip())
    return clean.strip("._")[:96] or "cover_letter"


def _unique_non_empty_texts(values: list[Any]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for raw in values:
        text = _coerce_text(raw)
        if not text:
            continue
        key = text.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(text)
    return out


def _search_profile_id_part(value: str) -> str:
    lowered = _coerce_text(value).lower()
    if not lowered:
        return "item"
    clean = _KEY_CHARS_PATTERN.sub("_", lowered).strip("._")
    return clean or "item"


def _next_profile_id(*, keyword: str, location: str, used: set[str]) -> str:
    base = f"{_search_profile_id_part(keyword)}_{_search_profile_id_part(location)}".strip("_")
    if not base:
        base = "search_profile"
    profile_id = base
    idx = 2
    while profile_id in used:
        profile_id = f"{base}_{idx}"
        idx += 1
    used.add(profile_id)
    return profile_id


def _assemble_search_profiles(cfg: dict[str, Any]) -> list[dict[str, str]]:
    # Backward compatibility first: if explicit profiles are provided, respect them.
    search_profiles_raw = cfg.get("search_profiles")
    search_profiles = [item for item in search_profiles_raw if isinstance(item, dict)] if isinstance(search_profiles_raw, list) else []
    out: list[dict[str, str]] = []
    used_ids: set[str] = set()
    for item in search_profiles:
        keyword = _coerce_text(item.get("keyword"))
        location = _coerce_text(item.get("location"))
        if not keyword or not location:
            continue
        profile_id = _coerce_text(item.get("profile_id")) or _next_profile_id(keyword=keyword, location=location, used=used_ids)
        if profile_id in used_ids:
            profile_id = _next_profile_id(keyword=keyword, location=location, used=used_ids)
        else:
            used_ids.add(profile_id)
        out.append({"profile_id": profile_id, "keyword": keyword, "location": location})
    if out:
        return out

    raw_locations = cfg.get("search_locations")
    raw_keywords = cfg.get("search_keywords")
    locations = _unique_non_empty_texts(raw_locations) if isinstance(raw_locations, list) else []
    keywords = _unique_non_empty_texts(raw_keywords) if isinstance(raw_keywords, list) else []

    # New layout: generate combinations from locations x keywords.
    if locations and keywords:
        used_ids: set[str] = set()
        out: list[dict[str, str]] = []
        for location in locations:
            for keyword in keywords:
                out.append(
                    {
                        "profile_id": _next_profile_id(keyword=keyword, location=location, used=used_ids),
                        "keyword": keyword,
                        "location": location,
                    }
                )
        return out

    return []


def _compact_file_segment(value: str, *, fallback: str, max_words: int = 4, max_chars: int = 28) -> str:
    text = _coerce_text(value)
    if not text or _is_not_found(text):
        text = fallback
    text = re.sub(r"\s+", " ", text).strip()
    text = re.split(r"\s(?:-|:|\|)\s", text, maxsplit=1)[0].strip() or text
    text = re.sub(r"[^A-Za-z0-9 ]+", " ", text)
    words = [w for w in text.split(" ") if w]
    if not words:
        words = [fallback]
    clipped = words[: max(1, max_words)]
    normalized = " ".join(clipped)
    if len(normalized) > max_chars:
        normalized = normalized[:max_chars].strip()
    if not normalized:
        normalized = fallback
    return normalized


def _next_available_local_docx_path(*, output_dir: Path, base_name: str, file_mode: str) -> Path:
    safe_base = _coerce_text(base_name) or "cover_letter"
    safe_base = safe_base.replace("/", " ").replace("\\", " ").strip()
    safe_base = re.sub(r"\s{2,}", " ", safe_base)
    output_dir.mkdir(parents=True, exist_ok=True)

    candidate = output_dir / f"{safe_base}.docx"
    if file_mode == "overwrite":
        return candidate
    if not candidate.exists():
        return candidate
    idx = 1
    while True:
        named = output_dir / f"{safe_base} - {idx}.docx"
        if not named.exists():
            return named
        idx += 1


def _repo_relative_path(path: Path) -> Path:
    try:
        return path.resolve().relative_to(_repo_root())
    except Exception:
        return Path("outputs") / "cover_letters" / "indeed_daily_search"


def _read_text_if_exists(path: Path) -> str:
    if not path.exists() or not path.is_file():
        return ""
    try:
        return path.read_text(encoding="utf-8").strip()
    except Exception:
        return ""


def _context_key_from_path(path: Path, prefix: str = "context") -> str:
    stem = _KEY_CHARS_PATTERN.sub("_", path.stem).strip("_").lower() or "item"
    return f"{prefix}_{stem}"


def _normalize_not_found(value: Any) -> str:
    text = _coerce_text(value)
    return text if text else NOT_FOUND_VALUE


def _is_not_found(value: Any) -> bool:
    return _normalize_not_found(value).strip().lower() == NOT_FOUND_VALUE.lower()


def _word_count(text: str) -> int:
    if not isinstance(text, str):
        return 0
    return len(_WORD_PATTERN.findall(text))


def _sanitize_cover_letter_text(text: str) -> str:
    cleaned = _coerce_text(text)
    if not cleaned:
        return ""
    cleaned = cleaned.replace("\u2014", ", ")
    cleaned = cleaned.replace("\u2013", ", ")
    # Remove hyphen punctuation in body text to honor no-dash style.
    cleaned = re.sub(r"(?<=\w)-(?=\w)", " ", cleaned)
    cleaned = re.sub(r"\s-\s", ", ", cleaned)
    cleaned = cleaned.replace("-", " ")
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned


def _normalize_role_title_for_cover_letter(raw_title: str) -> str:
    text = _coerce_text(raw_title)
    if not text:
        return "the role"
    cleaned = re.sub(r"\s+", " ", text).strip(" -|,:")
    cleaned = re.sub(r"\s*\([^)]*\)\s*$", "", cleaned).strip() or cleaned
    cleaned = re.sub(r"\s+[-:|/]\s+", " ", cleaned).strip() or cleaned

    # Remove level/seniority markers that make the title noisy.
    cleaned = re.sub(r"\b(?:L|Level)\s*\d+\b", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\b(?:I|II|III|IV|V|VI|VII|VIII|IX|X)\b", "", cleaned)
    cleaned = re.sub(r"\b(?:New\s+Grad|Early\s+Career|Campus\s+Hire)\b", "", cleaned, flags=re.IGNORECASE)

    # Remove broad department suffixes.
    cleaned = re.sub(r"\bInformation\s+Technology\b", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\bWeb\s*(?:&|and)\s*Mobile\s+Dev(?:elopment)?\b", "", cleaned, flags=re.IGNORECASE)

    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip(" ,:")
    if not cleaned:
        return "the role"

    alpha_chars = [ch for ch in cleaned if ch.isalpha()]
    uppercase_chars = [ch for ch in alpha_chars if ch.isupper()]
    if alpha_chars and (len(uppercase_chars) / len(alpha_chars)) >= 0.65:
        cleaned = cleaned.title()

    tokens = []
    for part in cleaned.split():
        lower = part.lower()
        if lower == "devops":
            tokens.append("DevOps")
            continue
        if lower in {"ai", "ml", "qa", "ui", "ux", "api", "sql", "sre", "is"}:
            tokens.append(lower.upper())
            continue
        tokens.append(part[:1].upper() + part[1:].lower())
    cleaned = " ".join(tokens)

    # Prefer concise role wording for common "Software Engineer <specialty>" titles.
    match = re.match(
        r"(?i)^Software Engineer\s+(DevOps|Backend|Frontend|Full Stack|Platform|Cloud|Security|Mobile|Automation)$",
        cleaned,
    )
    if match:
        cleaned = f"{match.group(1)} Software Engineer"

    return cleaned or "the role"


def _rewrite_role_title_mentions(text: str, *, raw_title: str, normalized_title: str) -> str:
    body = _coerce_text(text)
    raw = _coerce_text(raw_title)
    target = _coerce_text(normalized_title)
    if not body or not raw or not target:
        return body
    if raw.lower() == target.lower():
        return body
    return re.sub(re.escape(raw), target, body, flags=re.IGNORECASE)


@dataclass(frozen=True)
class CandidateContextBundle:
    base_context: dict[str, str]
    project_context: dict[str, str]


def _read_relative_context_file(path_text: str) -> tuple[str, str] | None:
    rel = _coerce_text(path_text)
    if not rel:
        return None
    path = _repo_root() / rel
    content = _read_text_if_exists(path)
    if not content:
        return None
    return _context_key_from_path(path), content


def _load_candidate_context_bundle(cfg: dict[str, Any]) -> CandidateContextBundle:
    root = _repo_root()
    configured_files = cfg.get("candidate_context_files")
    context_files = [item for item in configured_files if isinstance(item, str)] if isinstance(configured_files, list) else list(DEFAULT_CANDIDATE_CONTEXT_FILES)

    base_context: dict[str, str] = {}
    for item in context_files:
        loaded = _read_relative_context_file(item)
        if loaded is None:
            continue
        key, content = loaded
        base_context[key] = content

    configured_project_files = cfg.get("project_context_files")
    project_context_files = [item for item in configured_project_files if isinstance(item, str)] if isinstance(configured_project_files, list) else []
    if not project_context_files:
        default_project_dir = root / "context" / "more-about-human" / "projects"
        if default_project_dir.exists():
            for path in sorted(default_project_dir.glob("*.md")):
                if path.name.lower() == "project_index.md":
                    continue
                project_context_files.append(str(path.relative_to(root).as_posix()))

    project_context: dict[str, str] = {}
    for item in project_context_files:
        loaded = _read_relative_context_file(item)
        if loaded is None:
            continue
        _, content = loaded
        key = f"project_{Path(item).stem}"
        project_context[key] = content
    return CandidateContextBundle(base_context=base_context, project_context=project_context)


def _tokenize_for_scoring(text: str) -> set[str]:
    return set(_TOKEN_PATTERN.findall(text.lower()))


def _select_project_context_for_job(
    *,
    bundle: CandidateContextBundle,
    job_listing: dict[str, Any],
    job_detail: dict[str, Any],
    top_n: int,
    max_chars_per_project: int,
) -> dict[str, str]:
    if not bundle.project_context:
        return {}
    safe_top_n = max(0, int(top_n))
    if safe_top_n == 0:
        return {}

    job_text = json.dumps(job_listing, ensure_ascii=True) + "\n" + json.dumps(job_detail, ensure_ascii=True)
    job_tokens = _tokenize_for_scoring(job_text)

    ranked: list[tuple[int, str, str]] = []
    for key, content in bundle.project_context.items():
        score = len(job_tokens.intersection(_tokenize_for_scoring(content)))
        ranked.append((score, key, content))
    ranked.sort(key=lambda item: (-item[0], item[1]))

    selected = ranked[:safe_top_n]
    if not any(score > 0 for score, _, _ in selected):
        selected = ranked[:safe_top_n]

    out: dict[str, str] = {}
    safe_max = max(400, int(max_chars_per_project))
    for _, key, content in selected:
        text = content[:safe_max].strip()
        if text:
            out[key] = text
    return out


def _assemble_candidate_context_for_job(
    *,
    bundle: CandidateContextBundle,
    job_listing: dict[str, Any],
    job_detail: dict[str, Any],
    project_top_n: int,
    project_max_chars: int,
) -> dict[str, str]:
    context = dict(bundle.base_context)
    context.update(
        _select_project_context_for_job(
            bundle=bundle,
            job_listing=job_listing,
            job_detail=job_detail,
            top_n=project_top_n,
            max_chars_per_project=project_max_chars,
        )
    )
    return context


def _config_int(cfg: dict[str, Any], key: str, default: int, *, min_value: int = 1) -> int:
    value = cfg.get(key)
    if isinstance(value, int) and value >= min_value:
        return value
    return default


def _db_path_from_config() -> Path:
    cfg = load_config()
    central = get_central_service_config(cfg)
    raw = central.get("scheduler_db_path")
    return resolve_scheduler_db_path(str(raw) if isinstance(raw, str) else None)


def _connect_db(db_path: Path) -> sqlite3.Connection:
    conn = sqlite3.connect(db_path, timeout=10.0)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON;")
    conn.execute("PRAGMA busy_timeout = 10000;")
    return conn


def load_recent_seen_job_keys(
    *,
    task_id: str,
    provider: str = DEFAULT_PROVIDER,
    limit: int = DEFAULT_SEEN_LIMIT,
    db_path: Path | None = None,
) -> list[str]:
    safe_limit = max(1, int(limit))
    path = db_path or _db_path_from_config()
    with _connect_db(path) as conn:
        rows = conn.execute(
            """
            SELECT item_key
            FROM task_seen_items
            WHERE task_id = ? AND provider = ?
            ORDER BY first_seen_at DESC, item_key DESC
            LIMIT ?;
            """,
            (task_id, provider, safe_limit),
        ).fetchall()
    out: list[str] = []
    for row in rows:
        key = _coerce_text(row["item_key"])
        if key:
            out.append(key)
    return out


def _mark_job_seen(
    *,
    task_id: str,
    provider: str,
    job_key: str,
    metadata: dict[str, Any],
    db_path: Path,
) -> None:
    now = _iso_now()
    with _connect_db(db_path) as conn:
        conn.execute(
            """
            INSERT INTO task_seen_items(task_id, provider, item_key, metadata_json, first_seen_at, last_seen_at, seen_count)
            VALUES (?, ?, ?, ?, ?, ?, 1)
            ON CONFLICT(task_id, provider, item_key) DO UPDATE SET
                metadata_json = excluded.metadata_json,
                last_seen_at = excluded.last_seen_at,
                seen_count = task_seen_items.seen_count + 1;
            """,
            (
                task_id,
                provider,
                job_key,
                json.dumps(metadata, ensure_ascii=True),
                now,
                now,
            ),
        )


def _upsert_job_discovery(
    *,
    task_id: str,
    job_key: str,
    found_at: str,
    decision: str,
    db_path: Path,
) -> None:
    with _connect_db(db_path) as conn:
        conn.execute(
            """
            INSERT INTO job_discovery(task_id, job_key, found_at, decision, created_at)
            VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)
            ON CONFLICT(task_id, job_key) DO UPDATE SET
                found_at = excluded.found_at,
                decision = excluded.decision;
            """,
            (task_id, job_key, found_at, decision),
        )


def _upsert_task_state_snapshot(
    *,
    task_id: str,
    state_key: str,
    value: dict[str, Any],
    updated_by: str,
    db_path: Path,
) -> None:
    with _connect_db(db_path) as conn:
        conn.execute(
            """
            INSERT INTO task_state_kv(task_id, state_key, value_json, updated_at, updated_by)
            VALUES (?, ?, ?, ?, ?)
            ON CONFLICT(task_id, state_key) DO UPDATE SET
                value_json = excluded.value_json,
                updated_at = excluded.updated_at,
                updated_by = excluded.updated_by;
            """,
            (task_id, state_key, json.dumps(value, ensure_ascii=True), _iso_now(), updated_by),
        )


def _search_fraction(
    *,
    query_index: int,
    query_total: int,
    job_index: int,
    job_total: int,
    query_complete_no_jobs: bool = False,
) -> float:
    if query_total <= 0:
        return 0.0
    safe_query_index = max(1, min(query_total, int(query_index)))
    base = float(safe_query_index - 1) / float(query_total)
    if job_total > 0:
        within = (float(max(0, min(job_total, int(job_index)))) / float(job_total)) * (1.0 / float(query_total))
    else:
        within = (1.0 / float(query_total)) if query_complete_no_jobs else 0.0
    return max(0.0, min(1.0, base + within))


def _overall_fraction(*, search_fraction: float, processed_jobs: int, total_jobs: int) -> float:
    search_part = max(0.0, min(1.0, float(search_fraction)))
    if total_jobs <= 0:
        process_part = 0.0
    else:
        process_part = max(0.0, min(1.0, float(processed_jobs) / float(total_jobs)))
    return max(0.0, min(1.0, (SEARCH_PHASE_WEIGHT * search_part) + ((1.0 - SEARCH_PHASE_WEIGHT) * process_part)))


class ProgressReporter:
    def __init__(
        self,
        *,
        task_id: str,
        db_path: Path,
        callback: Callable[[dict[str, Any]], None] | None = None,
    ) -> None:
        self._task_id = task_id
        self._db_path = db_path
        self._callback = callback
        self._last_emit_key: tuple[str, int, int, int, int, int, int, int] | None = None

    def emit(self, payload: dict[str, Any]) -> None:
        status_line = _coerce_text(payload.get("status_line")) or "task progress update"
        stage = _coerce_text(payload.get("stage")) or "running"
        total_percent = float(payload.get("total_percent") or 0.0)
        search_fraction = float(payload.get("search_fraction") or 0.0)
        query_index = int(payload.get("query_index") or 0)
        query_total = int(payload.get("query_total") or 0)
        job_index = int(payload.get("job_index") or 0)
        job_total = int(payload.get("job_total") or 0)
        processed_jobs = int(payload.get("processed_jobs") or 0)
        total_jobs_for_processing = int(payload.get("total_jobs_for_processing") or 0)
        slot_update_seq = int(payload.get("slot_update_seq") or 0)
        emit_key = (
            stage,
            processed_jobs,
            total_jobs_for_processing,
            query_index,
            query_total,
            job_index,
            job_total,
            slot_update_seq,
        )
        if emit_key == self._last_emit_key:
            return
        self._last_emit_key = emit_key

        progress_row = {
            **payload,
            "search_percent": round(max(0.0, min(100.0, search_fraction * 100.0)), 1),
            "overall_percent": round(max(0.0, min(100.0, total_percent)), 1),
            "total_percent": round(total_percent, 1),
            "updated_at": _iso_now(),
        }
        try:
            _upsert_task_state_snapshot(
                task_id=self._task_id,
                state_key="live_progress",
                value=progress_row,
                updated_by="indeed_daily_search_task",
                db_path=self._db_path,
            )
        except Exception:
            pass
        if self._callback is not None:
            try:
                self._callback(progress_row)
            except Exception:
                pass

def _render_cover_letter_docx(
    *,
    output_path: Path,
    company_name: str,
    body_paragraphs: list[str],
    contact_email: str,
    linkedin_url: str,
    linkedin_label: str = "LinkedIn",
    date_line: str | None = None,
) -> None:
    from docx import Document  # type: ignore
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    def _append_hyperlink(*, paragraph: Any, url: str, text: str) -> None:
        if not _coerce_text(url) or not _coerce_text(text):
            return
        rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel_id)

        run = OxmlElement("w:r")
        run_props = OxmlElement("w:rPr")

        run_fonts = OxmlElement("w:rFonts")
        run_fonts.set(qn("w:ascii"), "Times New Roman")
        run_fonts.set(qn("w:hAnsi"), "Times New Roman")
        run_props.append(run_fonts)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        run_props.append(color)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_props.append(underline)

        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "24")
        run_props.append(size)
        size_cs = OxmlElement("w:szCs")
        size_cs.set(qn("w:val"), "24")
        run_props.append(size_cs)

        run.append(run_props)
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)

        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    # Name line
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_run = name_paragraph.add_run("Zubin Jha")
    name_run.bold = True
    name_run.font.name = "Times New Roman"
    name_run.font.size = Pt(21)

    # Address + phone line
    address_line = doc.add_paragraph()
    address_line.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    address_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = address_line.add_run("45 West Stafford Ave, Worthington, OH 43085\t(614)-653-3941")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Email + LinkedIn line
    contact_line = doc.add_paragraph()
    contact_line.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    contact_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    _append_hyperlink(paragraph=contact_line, url=f"mailto:{contact_email}", text=contact_email)
    spacer = contact_line.add_run("\t")
    spacer.font.name = "Times New Roman"
    spacer.font.size = Pt(12)
    _append_hyperlink(paragraph=contact_line, url=linkedin_url, text=linkedin_label)

    doc.add_paragraph("")

    date_text = date_line or _utc_now().strftime("%B %d, %Y")
    date_paragraph = doc.add_paragraph(date_text)
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in date_paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    salutation = doc.add_paragraph(f"Dear Hiring Manager at {company_name},")
    for run in salutation.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    for paragraph in body_paragraphs:
        p = doc.add_paragraph(paragraph.strip())
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    closing = doc.add_paragraph("Sincerely,")
    for run in closing.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    signature = doc.add_paragraph("Zubin Jha")
    for run in signature.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.save(str(output_path))


def _validate_decision_payload(raw: dict[str, Any]) -> tuple[bool, str]:
    decision = _coerce_text(raw.get("decision"))
    if decision not in ALLOWED_DECISIONS:
        return False, "decision must be one of Recommend Apply, Recommend Maybe, Skip"
    fit_score = raw.get("fit_score")
    if not isinstance(fit_score, int) or fit_score < 1 or fit_score > 10:
        return False, "fit_score must be integer 1-10"
    rationale = raw.get("rationale_short")
    if not isinstance(rationale, str) or not rationale.strip():
        return False, "rationale_short must be non-empty string"
    for list_key in ("reasons", "risks", "missing_requirements"):
        value = raw.get(list_key)
        if not isinstance(value, list):
            return False, f"{list_key} must be list"
        if any(not isinstance(item, str) or not item.strip() for item in value):
            return False, f"{list_key} must contain non-empty strings"
    return True, ""


def _validate_letter_payload(raw: dict[str, Any]) -> tuple[bool, str]:
    paragraphs = raw.get("paragraphs")
    if not isinstance(paragraphs, list):
        return False, "paragraphs must be a list"
    clean = [_sanitize_cover_letter_text(item) for item in paragraphs if isinstance(item, str) and _sanitize_cover_letter_text(item)]
    if len(clean) < 4:
        return False, "paragraphs must contain at least 4 non-empty entries"
    if len(clean) > 5:
        return False, "paragraphs must contain at most 5 entries"
    total_words = sum(_word_count(item) for item in clean)
    if total_words < 220:
        return False, "cover letter must be at least 220 words"
    for idx, item in enumerate(clean, start=1):
        if _word_count(item) < 35:
            return False, f"paragraph {idx} must be at least 35 words"
    return True, ""


def _validate_sheet_field_payload(raw: dict[str, Any]) -> tuple[bool, str]:
    expected = ["company", "job_title", "location", "pay_range", "job_link"]
    for key in expected:
        value = raw.get(key)
        if not isinstance(value, str):
            return False, f"{key} must be string"
    return True, ""


def _default_sheet_fields() -> dict[str, str]:
    return {
        "company": NOT_FOUND_VALUE,
        "job_title": NOT_FOUND_VALUE,
        "location": NOT_FOUND_VALUE,
        "pay_range": NOT_FOUND_VALUE,
        "job_link": NOT_FOUND_VALUE,
    }


def _candidate_context_text(candidate_context: dict[str, str]) -> str:
    return "\n\n".join([f"[{k}]\n{v}" for k, v in candidate_context.items() if _coerce_text(v)])


def _extract_sheet_fields_via_llm(
    *,
    model_alias: str,
    job_listing: dict[str, Any],
    job_detail: dict[str, Any],
    invalid_retry_limit: int,
) -> dict[str, Any]:
    listing_json = json.dumps(job_listing, ensure_ascii=True)
    detail_json = json.dumps(job_detail.get("job") if isinstance(job_detail.get("job"), dict) else {}, ensure_ascii=True)
    prompt = (
        "Extract canonical spreadsheet fields from this job listing/detail payload.\n"
        f"Return JSON only with keys company, job_title, location, pay_range, job_link.\n"
        f"When a value is missing or ambiguous return exact string: {NOT_FOUND_VALUE}\n"
        "Do not include markdown.\n\n"
        f"[JobListing]\n{listing_json}\n\n"
        f"[JobDetail]\n{detail_json}\n"
    )
    system = "You are a strict JSON extraction engine. Return only valid JSON."
    attempts = max(0, int(invalid_retry_limit)) + 1
    last_error = "sheet_field_extraction_failed"
    for attempt in range(1, attempts + 1):
        result = _llm_json_response(model_alias=model_alias, system_prompt=system, user_prompt=prompt)
        if not result.get("ok"):
            last_error = str(result.get("error") or "llm_error")
            if attempt < attempts:
                prompt += f"\n\nValidation error from prior attempt: {last_error}. Fix JSON strictly."
            continue
        payload = result["payload"]
        valid, reason = _validate_sheet_field_payload(payload)
        if valid:
            return {
                "ok": True,
                "fields": {
                    "company": _normalize_not_found(payload.get("company")),
                    "job_title": _normalize_not_found(payload.get("job_title")),
                    "location": _normalize_not_found(payload.get("location")),
                    "pay_range": _normalize_not_found(payload.get("pay_range")),
                    "job_link": _normalize_not_found(payload.get("job_link")),
                },
            }
        last_error = reason
        if attempt < attempts:
            prompt += f"\n\nValidation error from prior attempt: {reason}. Fix JSON strictly."
    return {"ok": False, "error": last_error, "fields": _default_sheet_fields()}


def _extract_first_json_object(text: str) -> dict[str, Any] | None:
    if not isinstance(text, str):
        return None
    raw = text.strip()
    if not raw:
        return None
    # Strip common fenced wrappers if present.
    raw = raw.replace("```json", "```")
    raw = raw.replace("```JSON", "```")
    raw = raw.strip("`").strip()
    # Try direct parse first.
    try:
        payload = json.loads(raw)
        if isinstance(payload, dict):
            return payload
    except Exception:
        pass

    # Best-effort balanced brace extraction.
    for start_idx, ch in enumerate(raw):
        if ch != "{":
            continue
        depth = 0
        for end_idx in range(start_idx, len(raw)):
            token = raw[end_idx]
            if token == "{":
                depth += 1
            elif token == "}":
                depth -= 1
                if depth == 0:
                    candidate = raw[start_idx : end_idx + 1]
                    try:
                        payload = json.loads(candidate)
                    except Exception:
                        break
                    if isinstance(payload, dict):
                        return payload
                    break
        # continue searching from next start_idx
    return None


def _llm_json_response(
    *,
    model_alias: str,
    system_prompt: str,
    user_prompt: str,
) -> dict[str, Any]:
    out = call_llm(
        model=model_alias,
        max_output_tokens=1200,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )
    if not out.get("ok"):
        return {"ok": False, "error": str(out.get("error") or "llm_error")}
    text = _coerce_text(out.get("text"))
    if not text:
        return {"ok": False, "error": "llm_empty_text"}
    payload = _extract_first_json_object(text)
    if not isinstance(payload, dict):
        return {"ok": False, "error": "invalid_json", "raw_text": text}
    return {"ok": True, "payload": payload, "raw_text": text}


def _append_sheet_row_with_retry(
    *,
    row: dict[str, Any],
    attempts: int,
    backoff_sec: float,
) -> dict[str, Any]:
    safe_attempts = max(1, int(attempts))
    safe_backoff = max(0.0, float(backoff_sec))
    last: dict[str, Any] = {"ok": False, "error": "unknown"}
    for idx in range(1, safe_attempts + 1):
        out = append_job_app_row(row=row)
        if bool(out.get("ok")):
            return {**out, "attempts_used": idx, "attempts_configured": safe_attempts}
        last = out if isinstance(out, dict) else {"ok": False, "error": "invalid_response"}
        if idx >= safe_attempts:
            break
        if safe_backoff > 0:
            sleep(safe_backoff * idx)
    return {**last, "attempts_used": safe_attempts, "attempts_configured": safe_attempts}


def _upload_cover_letter_with_retry(
    *,
    local_path: str,
    destination_path: str,
    destination_folder_id: str | None,
    filename: str,
    attempts: int,
    backoff_sec: float,
) -> dict[str, Any]:
    safe_attempts = max(1, int(attempts))
    safe_backoff = max(0.0, float(backoff_sec))
    last: dict[str, Any] = {"ok": False, "error": "unknown"}
    for idx in range(1, safe_attempts + 1):
        out = upload_file_to_google_drive(
            local_path=local_path,
            destination_path=destination_path,
            destination_folder_id=destination_folder_id,
            filename=filename,
        )
        if bool(out.get("ok")):
            return {**out, "attempts_used": idx, "attempts_configured": safe_attempts}
        last = out if isinstance(out, dict) else {"ok": False, "error": "invalid_response"}
        if idx >= safe_attempts:
            break
        if safe_backoff > 0:
            sleep(safe_backoff * idx)
    return {**last, "attempts_used": safe_attempts, "attempts_configured": safe_attempts}


def _evaluate_job(
    *,
    model_alias: str,
    job_listing: dict[str, Any],
    job_detail: dict[str, Any],
    candidate_context: dict[str, str],
    decision_rubric_text: str,
    invalid_retry_limit: int,
) -> dict[str, Any]:
    listing_json = json.dumps(job_listing, ensure_ascii=True)
    detail_json = json.dumps(job_detail.get("job") if isinstance(job_detail.get("job"), dict) else {}, ensure_ascii=True)
    context_text = _candidate_context_text(candidate_context)
    prompt = (
        "Decide whether this job should be Recommend Apply, Recommend Maybe, or Skip.\n"
        "Output JSON only with keys: decision, fit_score, rationale_short, reasons, risks, missing_requirements.\n"
        "Do not include markdown.\n\n"
        f"[DecisionRubric]\n{decision_rubric_text}\n\n"
        f"[CandidateContext]\n{context_text}\n\n"
        f"[JobListing]\n{listing_json}\n\n"
        f"[JobDetail]\n{detail_json}\n"
    )
    system = "You are a strict job application triage engine. Return valid JSON only."
    attempts = max(0, int(invalid_retry_limit)) + 1
    last_error = "decision_validation_failed"
    for attempt in range(1, attempts + 1):
        result = _llm_json_response(model_alias=model_alias, system_prompt=system, user_prompt=prompt)
        if not result.get("ok"):
            last_error = str(result.get("error") or "llm_error")
            if attempt < attempts:
                prompt += f"\n\nValidation error from prior attempt: {last_error}. Fix JSON strictly."
            continue
        payload = result["payload"]
        valid, reason = _validate_decision_payload(payload)
        if valid:
            return {"ok": True, "decision_payload": payload}
        last_error = reason
        if attempt < attempts:
            prompt += f"\n\nValidation error from prior attempt: {reason}. Fix JSON strictly."
    return {"ok": False, "error": last_error}


def _generate_cover_letter(
    *,
    model_alias: str,
    job_listing: dict[str, Any],
    job_detail: dict[str, Any],
    candidate_context: dict[str, str],
    style_spec_text: str,
    invalid_retry_limit: int,
) -> dict[str, Any]:
    raw_role_title = _extract_job_title(job_listing) or _extract_from_dict(
        job_detail.get("job") if isinstance(job_detail.get("job"), dict) else {},
        ["title", "jobTitle", "position"],
    )
    normalized_role_title = _normalize_role_title_for_cover_letter(raw_role_title)
    listing_json = json.dumps(job_listing, ensure_ascii=True)
    detail_json = json.dumps(job_detail.get("job") if isinstance(job_detail.get("job"), dict) else {}, ensure_ascii=True)
    context_text = _candidate_context_text(candidate_context)
    prompt = (
        "Write a tailored cover letter body that sounds like the candidate profile.\n"
        "Output JSON only with key paragraphs: array of exactly 4 paragraphs.\n"
        "Each paragraph should be between 45 and 95 words.\n"
        "Total word count should be between 220 and 380 words.\n"
        "No dash punctuation in body text.\n"
        "When mentioning the role title, use concise natural wording in title case.\n"
        "Do not copy long all-caps or department-heavy posting titles directly.\n"
        "Drop level/seniority suffixes like II, III, Level 2, and noisy department qualifiers.\n"
        "Prefer normalized role wording such as DevOps Software Engineer.\n"
        "Do not use markdown.\n\n"
        f"[RoleTitleRaw]\n{raw_role_title or 'Not Found'}\n\n"
        f"[RoleTitlePreferred]\n{normalized_role_title}\n\n"
        f"[StyleSpec]\n{style_spec_text}\n\n"
        f"[CandidateContext]\n{context_text}\n\n"
        f"[JobListing]\n{listing_json}\n\n"
        f"[JobDetail]\n{detail_json}\n"
    )
    system = "You write concise, specific cover letter paragraphs. Return valid JSON only."
    attempts = max(0, int(invalid_retry_limit)) + 1
    last_error = "letter_validation_failed"
    for attempt in range(1, attempts + 1):
        result = _llm_json_response(model_alias=model_alias, system_prompt=system, user_prompt=prompt)
        if not result.get("ok"):
            last_error = str(result.get("error") or "llm_error")
            if attempt < attempts:
                prompt += f"\n\nValidation error from prior attempt: {last_error}. Fix JSON strictly."
            continue
        payload = result["payload"]
        valid, reason = _validate_letter_payload(payload)
        if valid:
            paragraphs = [
                _sanitize_cover_letter_text(
                    _rewrite_role_title_mentions(
                        str(item),
                        raw_title=raw_role_title,
                        normalized_title=normalized_role_title,
                    )
                )
                for item in payload.get("paragraphs", [])
                if isinstance(item, str) and _sanitize_cover_letter_text(str(item))
            ]
            return {"ok": True, "paragraphs": paragraphs}
        last_error = reason
        if attempt < attempts:
            prompt += f"\n\nValidation error from prior attempt: {reason}. Fix JSON strictly."
    # Robust deterministic fallback so apply/maybe jobs still produce a cover letter.
    company = _extract_job_company(job_listing) or _extract_from_dict(
        job_detail.get("job") if isinstance(job_detail.get("job"), dict) else {},
        ["companyName", "company", "employer"],
    )
    title = normalized_role_title or "the role"
    company = company or "the company"
    paragraphs = [
        (
            f"I am excited to apply for {title} at {company}. I recently completed my Computer Science degree and I focus on building reliable software systems that solve practical problems. My background combines backend development, data focused engineering, and real product delivery. I am most engaged when I can turn messy requirements into clear technical plans and working software that teams can trust and improve over time."
        ),
        (
            "Across my projects, I have taken full ownership of systems from initial architecture through implementation, testing, and iteration. I have built API driven applications, relational database workflows, and analytics tools that required careful schema design, performance minded querying, and maintainable backend logic. This hands on work strengthened my ability to communicate tradeoffs, debug issues quickly, and deliver clean software under realistic constraints."
        ),
        (
            f"I am drawn to {company} because this role aligns with how I work best: building dependable systems, collaborating across functions, and continuously improving technical quality. I would bring a builder mindset, strong curiosity, and a clear focus on delivering useful outcomes for users and internal stakeholders. I am ready to contribute quickly while continuing to grow in the technologies and domain priorities that matter most to your team."
        ),
        (
            f"Thank you for considering my application for {title}. I would welcome the opportunity to discuss how my background in software engineering, data systems, and project ownership can support {company}. I am confident I can contribute with disciplined execution, thoughtful collaboration, and a strong commitment to quality from day one. I appreciate your time and I look forward to speaking with you."
        ),
    ]
    return {
        "ok": True,
        "paragraphs": paragraphs,
        "fallback_used": True,
        "fallback_reason": last_error,
    }


def _map_sheet_row(
    *,
    job_key: str,
    extracted_fields: dict[str, str],
    decision_payload: dict[str, Any],
    date_found_iso: str,
    job_url_fallback: str | None = None,
    cover_letter_link: str | None,
    note_suffix: str | None = None,
) -> dict[str, str]:
    company = _normalize_not_found(extracted_fields.get("company"))
    title = _normalize_not_found(extracted_fields.get("job_title"))
    location = _normalize_not_found(extracted_fields.get("location"))
    pay_range = _normalize_not_found(extracted_fields.get("pay_range"))
    job_url = _normalize_not_found(extracted_fields.get("job_link"))
    if _is_not_found(job_url) and _coerce_text(job_url_fallback):
        job_url = _coerce_text(job_url_fallback)
    decision = _coerce_text(decision_payload.get("decision")) or DECISION_SKIP
    fit_score = decision_payload.get("fit_score")
    rationale = _coerce_text(decision_payload.get("rationale_short"))
    notes = f"fit_score={fit_score}; rationale={rationale}"
    if note_suffix:
        notes = f"{notes}; {note_suffix}"
    return {
        "JobKey": job_key,
        "Company": company or "Unknown",
        "Job Title": title or "Unknown",
        "Location": location or "Unknown",
        "Date Found": date_found_iso,
        "Date Applied": "",
        "Status": decision,
        "Pay Range": pay_range,
        "Job Link": job_url,
        "Source": "Indeed",
        "Cover Letter": cover_letter_link or "",
        "Notes": "",
        "AI Notes": notes[:500],
    }


def _apply_deterministic_field_fallback(
    *,
    extracted_fields: dict[str, str],
    job_listing: dict[str, Any],
    job_detail: dict[str, Any],
    job_url: str,
) -> dict[str, str]:
    merged = dict(_default_sheet_fields())
    merged.update({k: _normalize_not_found(v) for k, v in (extracted_fields or {}).items() if isinstance(k, str)})
    detail_job = job_detail.get("job") if isinstance(job_detail.get("job"), dict) else {}

    deterministic_company = _extract_job_company(job_listing) or _extract_from_dict(detail_job, ["companyName", "company", "employer"])
    deterministic_title = _extract_job_title(job_listing) or _extract_from_dict(detail_job, ["title", "jobTitle", "position"])
    deterministic_location = _extract_job_location(job_listing) or _extract_from_dict(
        detail_job, ["location", "jobLocation", "formattedLocation", "cityState"]
    )
    deterministic_link = _coerce_text(job_url) or _extract_from_dict(detail_job, ["url", "jobUrl", "job_link", "jobLink", "link"])

    if _is_not_found(merged.get("company")) and deterministic_company:
        merged["company"] = deterministic_company
    if _is_not_found(merged.get("job_title")) and deterministic_title:
        merged["job_title"] = deterministic_title
    if _is_not_found(merged.get("location")) and deterministic_location:
        merged["location"] = deterministic_location
    if _is_not_found(merged.get("job_link")) and deterministic_link:
        merged["job_link"] = deterministic_link
    return {
        "company": _normalize_not_found(merged.get("company")),
        "job_title": _normalize_not_found(merged.get("job_title")),
        "location": _normalize_not_found(merged.get("location")),
        "pay_range": _normalize_not_found(merged.get("pay_range")),
        "job_link": _normalize_not_found(merged.get("job_link")),
    }


@dataclass
class SearchCandidate:
    job_key: str
    job_listing: dict[str, Any]
    search_profile_id: str
    query_index: int
    query_total: int
    result_index: int
    results_total: int
    keyword: str
    location: str
    is_new: bool


@dataclass
class DbWriteRequest:
    task_id: str
    job_key: str
    found_at: str
    decision: str
    done: Event
    result: dict[str, Any] | None = None


@dataclass
class SheetWriteRequest:
    row: dict[str, Any]
    attempts: int
    backoff_sec: float
    done: Event
    result: dict[str, Any] | None = None


@dataclass
class ProcessOutcome:
    candidate: SearchCandidate
    job_key: str
    decision: str
    outcome: str
    job_url: str | None
    error_reason: str | None
    cover_letter_local_path: str | None
    cover_letter_drive_file_id: str | None
    cover_letter_drive_folder_id: str | None
    count_deltas: dict[str, int]
    job_result: dict[str, Any]
    errors: list[str]


def _collect_new_candidates(
    *,
    task_id: str,
    db_path: Path,
    search_profiles: list[dict[str, Any]],
    seen_limit: int,
    provider: str,
    progress: ProgressReporter | None = None,
) -> tuple[list[SearchCandidate], dict[str, Any], list[str]]:
    errors: list[str] = []
    stats: dict[str, Any] = {
        "queries_total": len(search_profiles),
        "queries_ok": 0,
        "jobs_returned_total": 0,
        "jobs_filtered_seen": 0,
        "jobs_new_total": 0,
        "per_query": [],
    }
    initial_seen = set(load_recent_seen_job_keys(task_id=task_id, provider=provider, limit=seen_limit, db_path=db_path))
    in_run_seen = set(initial_seen)
    stats["initial_seen_loaded"] = len(initial_seen)
    discovered: list[SearchCandidate] = []

    total_queries = len(search_profiles)
    for query_idx, profile in enumerate(search_profiles, start=1):
        profile_id = _coerce_text(profile.get("profile_id")) or "search_profile"
        keyword = _coerce_text(profile.get("keyword"))
        location = _coerce_text(profile.get("location"))
        if progress is not None:
            search_frac = _search_fraction(
                query_index=query_idx,
                query_total=total_queries,
                job_index=0,
                job_total=0,
                query_complete_no_jobs=False,
            )
            overall_pct = _overall_fraction(search_fraction=search_frac, processed_jobs=0, total_jobs=0) * 100.0
            progress.emit(
                {
                    "stage": "search",
                    "query_index": query_idx,
                    "query_total": total_queries,
                    "query_keyword": keyword,
                    "query_location": location,
                    "job_index": 0,
                    "job_total": 0,
                    "search_fraction": round(search_frac, 4),
                    "processed_jobs": 0,
                    "total_jobs_for_processing": 0,
                    "total_percent": overall_pct,
                    "status_line": (
                        f"searching query {query_idx}/{total_queries} ({keyword}, {location}), "
                        f"fetching listings..."
                    ),
                }
            )
        if not keyword or not location:
            errors.append(f"search profile `{profile_id}` missing keyword/location")
            continue
        result = get_indeed_jobs(keyword=keyword, location=location)
        query_stats = {
            "profile_id": profile_id,
            "keyword": keyword,
            "location": location,
            "ok": bool(result.get("ok")),
            "returned": 0,
            "filtered_seen": 0,
            "kept_new": 0,
        }
        if not result.get("ok"):
            errors.append(f"search `{profile_id}` failed: {result.get('error')}")
            stats["per_query"].append(query_stats)
            continue
        stats["queries_ok"] += 1
        jobs = result.get("jobs") if isinstance(result.get("jobs"), list) else []
        query_stats["returned"] = len(jobs)
        stats["jobs_returned_total"] += len(jobs)
        if progress is not None and not jobs:
            search_frac = _search_fraction(
                query_index=query_idx,
                query_total=total_queries,
                job_index=0,
                job_total=0,
                query_complete_no_jobs=True,
            )
            overall_pct = _overall_fraction(search_fraction=search_frac, processed_jobs=0, total_jobs=0) * 100.0
            progress.emit(
                {
                    "stage": "search",
                    "query_index": query_idx,
                    "query_total": total_queries,
                    "query_keyword": keyword,
                    "query_location": location,
                    "job_index": 0,
                    "job_total": 0,
                    "search_fraction": round(search_frac, 4),
                    "processed_jobs": 0,
                    "total_jobs_for_processing": 0,
                    "total_percent": overall_pct,
                    "status_line": (
                        f"searching query {query_idx}/{total_queries} ({keyword}, {location}), no jobs returned"
                    ),
                }
            )
        if progress is not None and jobs:
            search_frac = _search_fraction(
                query_index=query_idx,
                query_total=total_queries,
                job_index=0,
                job_total=len(jobs),
                query_complete_no_jobs=False,
            )
            overall_pct = _overall_fraction(search_fraction=search_frac, processed_jobs=0, total_jobs=0) * 100.0
            progress.emit(
                {
                    "stage": "search",
                    "query_index": query_idx,
                    "query_total": total_queries,
                    "query_keyword": keyword,
                    "query_location": location,
                    "job_index": 0,
                    "job_total": len(jobs),
                    "search_fraction": round(search_frac, 4),
                    "processed_jobs": 0,
                    "total_jobs_for_processing": 0,
                    "total_percent": overall_pct,
                    "status_line": (
                        f"searching query {query_idx}/{total_queries} ({keyword}, {location}), "
                        f"received {len(jobs)} result(s)"
                    ),
                }
            )

        for job_idx, job in enumerate(jobs, start=1):
            if not isinstance(job, dict):
                continue
            job_key = _extract_job_key(job)
            if not job_key:
                continue
            is_new = job_key not in in_run_seen
            if not is_new:
                query_stats["filtered_seen"] += 1
                stats["jobs_filtered_seen"] += 1
            else:
                in_run_seen.add(job_key)
                query_stats["kept_new"] += 1
                stats["jobs_new_total"] += 1
                try:
                    _mark_job_seen(
                        task_id=task_id,
                        provider=provider,
                        job_key=job_key,
                        metadata={
                            "search_profile_id": profile_id,
                            "keyword": keyword,
                            "location": location,
                            "job_url": _extract_job_url(job),
                            "job_title": _extract_job_title(job),
                        },
                        db_path=db_path,
                    )
                except Exception as exc:  # pragma: no cover - defensive runtime guard
                    errors.append(f"mark_seen failed for `{job_key}`: {exc}")
            discovered.append(
                SearchCandidate(
                    job_key=job_key,
                    job_listing=job,
                    search_profile_id=profile_id,
                    query_index=query_idx,
                    query_total=total_queries,
                    result_index=job_idx,
                    results_total=len(jobs),
                    keyword=keyword,
                    location=location,
                    is_new=is_new,
                )
            )
        stats["per_query"].append(query_stats)
    return discovered, stats, errors


def _db_writer_loop(*, db_path: Path, queue: Queue[DbWriteRequest | object], stop_token: object) -> None:
    while True:
        item = queue.get()
        try:
            if item is stop_token:
                return
            assert isinstance(item, DbWriteRequest)
            try:
                _upsert_job_discovery(
                    task_id=item.task_id,
                    job_key=item.job_key,
                    found_at=item.found_at,
                    decision=item.decision,
                    db_path=db_path,
                )
                item.result = {"ok": True}
            except Exception as exc:  # pragma: no cover - runtime guard
                item.result = {"ok": False, "error": str(exc)}
            finally:
                item.done.set()
        finally:
            queue.task_done()


def _sheet_writer_loop(*, queue: Queue[SheetWriteRequest | object], stop_token: object) -> None:
    while True:
        item = queue.get()
        try:
            if item is stop_token:
                return
            assert isinstance(item, SheetWriteRequest)
            try:
                item.result = _append_sheet_row_with_retry(
                    row=item.row,
                    attempts=item.attempts,
                    backoff_sec=item.backoff_sec,
                )
            except Exception as exc:  # pragma: no cover - runtime guard
                item.result = {"ok": False, "error": str(exc), "source": "sheet_writer_exception"}
            finally:
                item.done.set()
        finally:
            queue.task_done()


def run_pipeline(
    *,
    task_id: str,
    payload: dict[str, Any],
    local_config: dict[str, Any],
    resources_dir: Path,
    progress_callback: Callable[[dict[str, Any]], None] | None = None,
) -> dict[str, Any]:
    cfg = local_config if isinstance(local_config, dict) else {}
    search_profiles = _assemble_search_profiles(cfg)
    if not search_profiles:
        return {
            "ok": False,
            "error": "task_config requires either both search_locations+search_keywords or search_profiles",
        }

    seen_limit = _config_int(cfg, "seen_ids_limit", DEFAULT_SEEN_LIMIT, min_value=1)
    provider = _coerce_text(cfg.get("seen_provider")) or DEFAULT_PROVIDER
    extraction_model_alias = _coerce_text(cfg.get("extraction_model_alias")) or DEFAULT_EXTRACTION_MODEL
    decision_model_alias = _coerce_text(cfg.get("decision_model_alias")) or DEFAULT_DECISION_MODEL
    cover_letter_model_alias = _coerce_text(cfg.get("cover_letter_model_alias")) or DEFAULT_COVER_LETTER_MODEL
    invalid_retry_limit = _config_int(cfg, "invalid_schema_retry_limit", 1, min_value=0)
    destination_path = _coerce_text(cfg.get("cover_letter_destination_path")) or DEFAULT_COVER_LETTER_DESTINATION_PATH
    destination_folder_id = _coerce_text(cfg.get("cover_letter_destination_folder_id")) or None
    file_mode = _coerce_text(cfg.get("cover_letter_file_mode")).lower() or DEFAULT_FILE_MODE
    contact_email = _coerce_text(cfg.get("cover_letter_contact_email")) or DEFAULT_COVER_LETTER_CONTACT_EMAIL
    linkedin_url = _coerce_text(cfg.get("cover_letter_linkedin_url")) or DEFAULT_COVER_LETTER_LINKEDIN_URL
    linkedin_label = _coerce_text(cfg.get("cover_letter_linkedin_label")) or DEFAULT_COVER_LETTER_LINKEDIN_LABEL
    cover_letter_upload_retry_attempts = _config_int(
        cfg,
        "cover_letter_upload_retry_attempts",
        DEFAULT_COVER_LETTER_UPLOAD_RETRY_ATTEMPTS,
        min_value=1,
    )
    cover_letter_upload_retry_backoff_sec = (
        float(cfg.get("cover_letter_upload_retry_backoff_sec"))
        if isinstance(cfg.get("cover_letter_upload_retry_backoff_sec"), (int, float))
        else DEFAULT_COVER_LETTER_UPLOAD_RETRY_BACKOFF_SEC
    )
    sheet_retry_attempts = _config_int(cfg, "sheet_retry_attempts", DEFAULT_SHEET_RETRY_ATTEMPTS, min_value=1)
    sheet_retry_backoff_sec = float(cfg.get("sheet_retry_backoff_sec")) if isinstance(cfg.get("sheet_retry_backoff_sec"), (int, float)) else DEFAULT_SHEET_RETRY_BACKOFF_SEC
    project_context_top_n = _config_int(cfg, "project_context_top_n", DEFAULT_PROJECT_CONTEXT_TOP_N, min_value=0)
    project_context_max_chars = _config_int(cfg, "project_context_max_chars", DEFAULT_PROJECT_CONTEXT_MAX_CHARS, min_value=300)
    process_workers = min(
        MAX_PROCESS_WORKERS,
        _config_int(cfg, "process_workers", DEFAULT_PROCESS_WORKERS, min_value=1),
    )
    db_queue_maxsize = _config_int(cfg, "db_queue_maxsize", DEFAULT_DB_QUEUE_MAXSIZE, min_value=1)
    sheet_queue_maxsize = _config_int(cfg, "sheet_queue_maxsize", DEFAULT_SHEET_QUEUE_MAXSIZE, min_value=1)
    if file_mode not in {"overwrite", "versioned"}:
        file_mode = DEFAULT_FILE_MODE

    db_path = _db_path_from_config()
    progress = ProgressReporter(task_id=task_id, db_path=db_path, callback=progress_callback)
    progress.emit(
        {
            "stage": "starting",
            "query_index": 0,
            "query_total": len(search_profiles),
            "job_index": 0,
            "job_total": 0,
            "search_fraction": 0.0,
            "processed_jobs": 0,
            "total_jobs_for_processing": 0,
            "total_percent": 0.0,
            "status_line": f"starting indeed_daily_search with {len(search_profiles)} query combination(s)...",
        }
    )
    discovered, search_stats, errors = _collect_new_candidates(
        task_id=task_id,
        db_path=db_path,
        search_profiles=search_profiles,
        seen_limit=seen_limit,
        provider=provider,
        progress=progress,
    )

    candidate_context_bundle = _load_candidate_context_bundle(cfg)
    decision_rubric_text = _read_text_if_exists(resources_dir / "assets" / "decision_rubric.md")
    style_spec_text = _read_text_if_exists(resources_dir / "assets" / "cover_letter_style_spec.md")
    if not decision_rubric_text:
        errors.append("missing decision_rubric.md; using built-in fallback rubric")
        decision_rubric_text = (
            "Recommend Apply for clear alignment and interview viability.\n"
            "Recommend Maybe for partial alignment with meaningful upside.\n"
            "Skip for clear mismatch or seniority gap."
        )
    if not style_spec_text:
        errors.append("missing cover_letter_style_spec.md; using built-in style fallback")
        style_spec_text = "Times New Roman, clear concrete language, no em dash punctuation."

    counts = {
        "searched": int(search_stats.get("queries_total") or 0),
        "new_jobs": int(search_stats.get("jobs_new_total") or 0),
        "seen_filtered": int(search_stats.get("jobs_filtered_seen") or 0),
        "recommended_apply": 0,
        "recommended_maybe": 0,
        "skipped": 0,
        "extraction_errors": 0,
        "decision_errors": 0,
        "cover_letter_errors": 0,
        "upload_errors": 0,
        "sheet_rows_written": 0,
        "sheet_rows_deduped": 0,
    }
    job_results: list[dict[str, Any]] = []
    found_at_iso = _local_today_iso()
    total_jobs_to_process = len(discovered)
    processed_jobs = 0
    worker_step_total = 4
    slot_update_seq = 0
    slot_lock = Lock()

    def _idle_slot(slot: int) -> dict[str, Any]:
        return {
            "slot": slot,
            "state": "idle",
            "step_key": "idle",
            "step_label": "Idle",
            "step_index": 0,
            "step_total": worker_step_total,
            "job_key": "",
            "query_index": 0,
            "query_total": 0,
            "result_index": 0,
            "result_total": 0,
            "query_keyword": "",
            "query_location": "",
            "summary": "idle",
            "updated_at": _iso_now(),
        }

    worker_slots: dict[int, dict[str, Any]] = {slot: _idle_slot(slot) for slot in range(1, process_workers + 1)}

    def _worker_slots_snapshot() -> tuple[int, list[dict[str, Any]]]:
        with slot_lock:
            slots = [dict(state) for _slot, state in sorted(worker_slots.items(), key=lambda item: item[0])]
            seq = int(slot_update_seq)
        return seq, slots

    def _set_slot_state(
        *,
        slot: int,
        state: str,
        step_key: str,
        step_label: str,
        step_index: int,
        candidate: SearchCandidate | None = None,
        job_key: str | None = None,
        emit_progress: bool = True,
    ) -> None:
        nonlocal slot_update_seq
        with slot_lock:
            base = dict(worker_slots.get(slot) or _idle_slot(slot))
            if candidate is not None:
                base.update(
                    {
                        "query_index": candidate.query_index,
                        "query_total": candidate.query_total,
                        "result_index": candidate.result_index,
                        "result_total": candidate.results_total,
                        "query_keyword": candidate.keyword,
                        "query_location": candidate.location,
                    }
                )
            if job_key is not None:
                base["job_key"] = str(job_key)
            clean_label = _coerce_text(step_label) or "running"
            base.update(
                {
                    "slot": slot,
                    "state": _coerce_text(state) or "running",
                    "step_key": _coerce_text(step_key) or "running",
                    "step_label": clean_label,
                    "step_index": max(0, min(worker_step_total, int(step_index))),
                    "step_total": worker_step_total,
                    "summary": f"{clean_label} ({base.get('job_key') or '-'})",
                    "updated_at": _iso_now(),
                }
            )
            worker_slots[slot] = base
            slot_update_seq += 1
            seq = int(slot_update_seq)
            slots = [dict(state_row) for _slot, state_row in sorted(worker_slots.items(), key=lambda item: item[0])]

        if not emit_progress:
            return

        progress.emit(
            {
                "stage": "process",
                "query_index": int(base.get("query_index") or 0),
                "query_total": int(base.get("query_total") or 0),
                "query_keyword": _coerce_text(base.get("query_keyword")),
                "query_location": _coerce_text(base.get("query_location")),
                "job_index": int(base.get("result_index") or 0),
                "job_total": int(base.get("result_total") or 0),
                "job_key": _coerce_text(base.get("job_key")),
                "search_fraction": 1.0,
                "processed_jobs": processed_jobs,
                "total_jobs_for_processing": total_jobs_to_process,
                "worker_slots": slots,
                "slot_update_seq": seq,
                "total_percent": _overall_fraction(search_fraction=1.0, processed_jobs=processed_jobs, total_jobs=total_jobs_to_process)
                * 100.0,
                "status_line": f"slot {slot}: {_coerce_text(base.get('summary'))}",
            }
        )

    progress.emit(
        {
            "stage": "process",
            "query_index": 0,
            "query_total": int(search_stats.get("queries_total") or 0),
            "job_index": 0,
            "job_total": total_jobs_to_process,
            "search_fraction": 0.0,
            "processed_jobs": 0,
            "total_jobs_for_processing": total_jobs_to_process,
            "worker_slots": _worker_slots_snapshot()[1],
            "slot_update_seq": slot_update_seq,
            "total_percent": _overall_fraction(search_fraction=1.0, processed_jobs=0, total_jobs=total_jobs_to_process) * 100.0,
            "status_line": f"search phase complete, processing {total_jobs_to_process} result(s)...",
        }
    )

    def _emit_job_result(
        *,
        candidate: SearchCandidate,
        job_key: str,
        decision: str,
        outcome: str,
        job_url: str | None = None,
        error_reason: str | None = None,
        cover_letter_local_path: str | None = None,
        cover_letter_drive_file_id: str | None = None,
        cover_letter_drive_folder_id: str | None = None,
    ) -> None:
        status_line = (
            f"query {candidate.query_index}/{candidate.query_total} "
            f"({candidate.keyword}, {candidate.location}), "
            f"result {candidate.result_index}/{candidate.results_total} ({job_key}) "
            f"decision={decision} outcome={outcome}"
        )
        if job_url:
            status_line += f" job_url={job_url}"
        if error_reason:
            status_line += f" error={error_reason}"
        if cover_letter_local_path:
            status_line += f" cover_letter_local_path={cover_letter_local_path}"
        seq, slots = _worker_slots_snapshot()
        progress.emit(
            {
                "stage": "process_result",
                "query_index": candidate.query_index,
                "query_total": candidate.query_total,
                "query_keyword": candidate.keyword,
                "query_location": candidate.location,
                "job_index": candidate.result_index,
                "job_total": candidate.results_total,
                "job_key": job_key,
                "job_url": job_url,
                "decision": decision,
                "outcome": outcome,
                "error_reason": error_reason,
                "cover_letter_local_path": cover_letter_local_path,
                "cover_letter_drive_file_id": cover_letter_drive_file_id,
                "cover_letter_drive_folder_id": cover_letter_drive_folder_id,
                "search_fraction": 1.0,
                "processed_jobs": processed_jobs,
                "total_jobs_for_processing": total_jobs_to_process,
                "worker_slots": slots,
                "slot_update_seq": seq,
                "total_percent": _overall_fraction(
                    search_fraction=1.0,
                    processed_jobs=processed_jobs,
                    total_jobs=total_jobs_to_process,
                )
                * 100.0,
                "status_line": status_line,
            }
        )
    db_stop_token = object()
    sheet_stop_token = object()
    db_write_queue: Queue[DbWriteRequest | object] = Queue(maxsize=max(1, db_queue_maxsize))
    sheet_write_queue: Queue[SheetWriteRequest | object] = Queue(maxsize=max(1, sheet_queue_maxsize))
    db_writer = Thread(
        target=_db_writer_loop,
        kwargs={"db_path": db_path, "queue": db_write_queue, "stop_token": db_stop_token},
        daemon=True,
        name="indeed_db_writer",
    )
    sheet_writer = Thread(
        target=_sheet_writer_loop,
        kwargs={"queue": sheet_write_queue, "stop_token": sheet_stop_token},
        daemon=True,
        name="indeed_sheet_writer",
    )
    db_writer.start()
    sheet_writer.start()

    def _enqueue_db_write(*, job_key: str, decision: str, out_errors: list[str]) -> None:
        req = DbWriteRequest(task_id=task_id, job_key=job_key, found_at=found_at_iso, decision=decision, done=Event())
        db_write_queue.put(req)
        req.done.wait()
        result = req.result if isinstance(req.result, dict) else {"ok": False, "error": "unknown_db_writer_error"}
        if not bool(result.get("ok")):
            out_errors.append(f"job_discovery write failed for `{job_key}`: {_coerce_text(result.get('error')) or 'unknown'}")

    def _enqueue_sheet_write(*, row: dict[str, Any]) -> dict[str, Any]:
        req = SheetWriteRequest(
            row=row,
            attempts=sheet_retry_attempts,
            backoff_sec=sheet_retry_backoff_sec,
            done=Event(),
        )
        sheet_write_queue.put(req)
        req.done.wait()
        if isinstance(req.result, dict):
            return req.result
        return {"ok": False, "error": "unknown_sheet_writer_error", "source": "sheet_writer"}

    def _process_candidate(candidate: SearchCandidate, *, slot: int) -> ProcessOutcome:
        listing = candidate.job_listing
        job_key = candidate.job_key
        job_url = _extract_job_url(listing)
        deltas = {
            "recommended_apply": 0,
            "recommended_maybe": 0,
            "skipped": 0,
            "extraction_errors": 0,
            "decision_errors": 0,
            "cover_letter_errors": 0,
            "upload_errors": 0,
            "sheet_rows_written": 0,
            "sheet_rows_deduped": 0,
        }
        local_errors: list[str] = []
        if not candidate.is_new:
            return ProcessOutcome(
                candidate=candidate,
                job_key=job_key,
                decision="SeenSkip",
                outcome="already_seen",
                job_url=job_url,
                error_reason=None,
                cover_letter_local_path=None,
                cover_letter_drive_file_id=None,
                cover_letter_drive_folder_id=None,
                count_deltas=deltas,
                job_result={"job_key": job_key, "decision": None, "status": "seen_skip"},
                errors=local_errors,
            )
        if not job_url:
            deltas["skipped"] += 1
            deltas["decision_errors"] += 1
            local_errors.append(f"missing job url for `{job_key}`")
            return ProcessOutcome(
                candidate=candidate,
                job_key=job_key,
                decision=DECISION_SKIP,
                outcome="missing_job_url",
                job_url=None,
                error_reason="missing_job_url",
                cover_letter_local_path=None,
                cover_letter_drive_file_id=None,
                cover_letter_drive_folder_id=None,
                count_deltas=deltas,
                job_result={"job_key": job_key, "decision": DECISION_SKIP, "status": "missing_job_url"},
                errors=local_errors,
            )

        detail = get_indeed_job_detail(url=job_url)
        if not detail.get("ok"):
            deltas["skipped"] += 1
            deltas["decision_errors"] += 1
            detail_err = _coerce_text(detail.get("error")) or "detail_fetch_error"
            local_errors.append(f"detail fetch failed for `{job_key}`: {detail_err}")
            return ProcessOutcome(
                candidate=candidate,
                job_key=job_key,
                decision=DECISION_SKIP,
                outcome="detail_fetch_error",
                job_url=job_url,
                error_reason=detail_err,
                cover_letter_local_path=None,
                cover_letter_drive_file_id=None,
                cover_letter_drive_folder_id=None,
                count_deltas=deltas,
                job_result={"job_key": job_key, "decision": DECISION_SKIP, "status": "detail_fetch_error"},
                errors=local_errors,
            )

        candidate_context = _assemble_candidate_context_for_job(
            bundle=candidate_context_bundle,
            job_listing=listing,
            job_detail=detail,
            project_top_n=project_context_top_n,
            project_max_chars=project_context_max_chars,
        )
        sheet_extract = _extract_sheet_fields_via_llm(
            model_alias=extraction_model_alias,
            job_listing=listing,
            job_detail=detail,
            invalid_retry_limit=invalid_retry_limit,
        )
        extracted_fields = sheet_extract.get("fields") if isinstance(sheet_extract.get("fields"), dict) else _default_sheet_fields()
        extracted_fields = _apply_deterministic_field_fallback(
            extracted_fields=extracted_fields,
            job_listing=listing,
            job_detail=detail,
            job_url=job_url,
        )
        if not sheet_extract.get("ok"):
            deltas["extraction_errors"] += 1
            local_errors.append(f"field extraction failed for `{job_key}`: {sheet_extract.get('error')}")

        _set_slot_state(
            slot=slot,
            state="running",
            step_key="decide",
            step_label="Decide",
            step_index=2,
            candidate=candidate,
            job_key=job_key,
            emit_progress=True,
        )

        decision_out = _evaluate_job(
            model_alias=decision_model_alias,
            job_listing=listing,
            job_detail=detail,
            candidate_context=candidate_context,
            decision_rubric_text=decision_rubric_text,
            invalid_retry_limit=invalid_retry_limit,
        )
        if not decision_out.get("ok"):
            deltas["skipped"] += 1
            deltas["decision_errors"] += 1
            decision_err = _coerce_text(decision_out.get("error")) or "decision_failed"
            local_errors.append(f"decision failed for `{job_key}`: {decision_err}")
            _set_slot_state(
                slot=slot,
                state="running",
                step_key="persist",
                step_label="Persist",
                step_index=4,
                candidate=candidate,
                job_key=job_key,
                emit_progress=True,
            )
            _enqueue_db_write(job_key=job_key, decision=DECISION_SKIP, out_errors=local_errors)
            return ProcessOutcome(
                candidate=candidate,
                job_key=job_key,
                decision=DECISION_SKIP,
                outcome="decision_error",
                job_url=job_url,
                error_reason=decision_err,
                cover_letter_local_path=None,
                cover_letter_drive_file_id=None,
                cover_letter_drive_folder_id=None,
                count_deltas=deltas,
                job_result={"job_key": job_key, "decision": DECISION_SKIP, "status": "decision_error", "error": decision_err},
                errors=local_errors,
            )

        decision_payload = decision_out["decision_payload"]
        decision = str(decision_payload["decision"])
        if decision == DECISION_RECOMMEND_APPLY:
            deltas["recommended_apply"] += 1
        elif decision == DECISION_RECOMMEND_MAYBE:
            deltas["recommended_maybe"] += 1
        else:
            deltas["skipped"] += 1

        if decision == DECISION_SKIP:
            _set_slot_state(
                slot=slot,
                state="running",
                step_key="cover_letter",
                step_label="Skip cover letter",
                step_index=3,
                candidate=candidate,
                job_key=job_key,
                emit_progress=True,
            )
            _set_slot_state(
                slot=slot,
                state="running",
                step_key="persist",
                step_label="Persist",
                step_index=4,
                candidate=candidate,
                job_key=job_key,
                emit_progress=True,
            )
            _enqueue_db_write(job_key=job_key, decision=decision, out_errors=local_errors)
            return ProcessOutcome(
                candidate=candidate,
                job_key=job_key,
                decision=decision,
                outcome="skipped",
                job_url=job_url,
                error_reason=None,
                cover_letter_local_path=None,
                cover_letter_drive_file_id=None,
                cover_letter_drive_folder_id=None,
                count_deltas=deltas,
                job_result={"job_key": job_key, "decision": decision, "status": "skipped"},
                errors=local_errors,
            )

        note_suffix_parts: list[str] = []
        cover_link: str | None = None
        cover_letter_local_path: str | None = None
        cover_letter_drive_file_id: str | None = None
        cover_letter_drive_folder_id: str | None = None

        _set_slot_state(
            slot=slot,
            state="running",
            step_key="cover_letter",
            step_label="Write cover letter",
            step_index=3,
            candidate=candidate,
            job_key=job_key,
            emit_progress=True,
        )

        letter_out = _generate_cover_letter(
            model_alias=cover_letter_model_alias,
            job_listing=listing,
            job_detail=detail,
            candidate_context=candidate_context,
            style_spec_text=style_spec_text,
            invalid_retry_limit=invalid_retry_limit,
        )
        if not letter_out.get("ok"):
            deltas["cover_letter_errors"] += 1
            err = _coerce_text(letter_out.get("error")) or "unknown"
            local_errors.append(f"cover letter generation failed for `{job_key}`: {err}")
            note_suffix_parts.append(f"cover_letter_error={err}")
        else:
            if bool(letter_out.get("fallback_used")):
                fallback_reason = _coerce_text(letter_out.get("fallback_reason")) or "fallback_used"
                note_suffix_parts.append(f"cover_letter_fallback={fallback_reason}")
            company_name = _coerce_text(extracted_fields.get("company"))
            if not company_name or company_name.lower() == NOT_FOUND_VALUE.lower():
                company_name = "the company"
            role_name = (
                _coerce_text(extracted_fields.get("job_title"))
                or _normalize_role_title_for_cover_letter(_extract_job_title(listing))
                or "Role"
            )
            found_date = _coerce_text(found_at_iso)[:10] or _utc_now().strftime("%Y-%m-%d")
            company_segment = _compact_file_segment(company_name, fallback="Company")
            role_segment = _compact_file_segment(role_name, fallback="Role")
            base_file_name = f"{found_date} - {company_segment} - {role_segment}"
            output_dir = _repo_root() / _repo_relative_path(resources_dir) / "state" / "cover_letters"
            absolute_output_path = _next_available_local_docx_path(
                output_dir=output_dir,
                base_name=base_file_name,
                file_mode=file_mode,
            )
            relative_output_path = absolute_output_path.relative_to(_repo_root())
            try:
                _render_cover_letter_docx(
                    output_path=absolute_output_path,
                    company_name=company_name,
                    body_paragraphs=letter_out["paragraphs"],
                    contact_email=contact_email,
                    linkedin_url=linkedin_url,
                    linkedin_label=linkedin_label,
                )
            except Exception as exc:
                deltas["cover_letter_errors"] += 1
                local_errors.append(f"cover letter render failed for `{job_key}`: {exc}")
                note_suffix_parts.append("cover_letter_error=render_failed")
            else:
                cover_letter_local_path = str(relative_output_path)
                upload_out = _upload_cover_letter_with_retry(
                    local_path=str(relative_output_path),
                    destination_path=destination_path,
                    destination_folder_id=destination_folder_id,
                    filename=absolute_output_path.name,
                    attempts=cover_letter_upload_retry_attempts,
                    backoff_sec=cover_letter_upload_retry_backoff_sec,
                )
                if not upload_out.get("ok"):
                    deltas["upload_errors"] += 1
                    err = _coerce_text(upload_out.get("error")) or "unknown"
                    src = _coerce_text(upload_out.get("source")) or "google_drive_upload_error"
                    local_errors.append(f"cover letter upload failed for `{job_key}`: {src}: {err}")
                    note_suffix_parts.append(f"cover_letter_error=upload_failed:{src}:{err}")
                else:
                    cover_link = _coerce_text(upload_out.get("web_view_link"))
                    cover_letter_drive_file_id = _coerce_text(upload_out.get("drive_file_id")) or None
                    cover_letter_drive_folder_id = (
                        _coerce_text(upload_out.get("destination_folder_id"))
                        or _coerce_text(upload_out.get("drive_folder_id"))
                        or None
                    )
                    if not cover_link:
                        drive_file_id = _coerce_text(upload_out.get("drive_file_id"))
                        if drive_file_id:
                            cover_link = f"https://drive.google.com/file/d/{drive_file_id}/view"
                        else:
                            note_suffix_parts.append("cover_letter_error=upload_missing_link")
                    if destination_folder_id and cover_letter_drive_folder_id and cover_letter_drive_folder_id != destination_folder_id:
                        mismatch = f"{cover_letter_drive_folder_id}!={destination_folder_id}"
                        local_errors.append(f"cover letter upload folder mismatch for `{job_key}`: {mismatch}")
                        note_suffix_parts.append(f"cover_letter_error=folder_mismatch:{mismatch}")

        _set_slot_state(
            slot=slot,
            state="running",
            step_key="persist",
            step_label="Persist",
            step_index=4,
            candidate=candidate,
            job_key=job_key,
            emit_progress=True,
        )
        _enqueue_db_write(job_key=job_key, decision=decision, out_errors=local_errors)

        row = _map_sheet_row(
            job_key=job_key,
            extracted_fields=extracted_fields,
            decision_payload=decision_payload,
            date_found_iso=found_at_iso,
            job_url_fallback=job_url,
            cover_letter_link=cover_link,
            note_suffix="; ".join(note_suffix_parts) if note_suffix_parts else None,
        )
        if _is_not_found(row.get("Job Link")):
            deltas["upload_errors"] += 1
            local_errors.append(f"sheet row missing job link for `{job_key}`; row skipped")
            return ProcessOutcome(
                candidate=candidate,
                job_key=job_key,
                decision=decision,
                outcome="missing_job_link",
                job_url=job_url,
                error_reason="missing_job_link",
                cover_letter_local_path=cover_letter_local_path,
                cover_letter_drive_file_id=cover_letter_drive_file_id,
                cover_letter_drive_folder_id=cover_letter_drive_folder_id,
                count_deltas=deltas,
                job_result={"job_key": job_key, "decision": decision, "status": "missing_job_link"},
                errors=local_errors,
            )

        append_out = _enqueue_sheet_write(row=row)
        if append_out.get("ok"):
            deltas["sheet_rows_written"] += 1
            status = "uploaded" if cover_link else "uploaded_without_cover_letter"
            return ProcessOutcome(
                candidate=candidate,
                job_key=job_key,
                decision=decision,
                outcome=status,
                job_url=job_url,
                error_reason=None,
                cover_letter_local_path=cover_letter_local_path,
                cover_letter_drive_file_id=cover_letter_drive_file_id,
                cover_letter_drive_folder_id=cover_letter_drive_folder_id,
                count_deltas=deltas,
                job_result={"job_key": job_key, "decision": decision, "status": status, "cover_letter": cover_link},
                errors=local_errors,
            )

        err_text = _coerce_text(append_out.get("error"))
        if "Duplicate JobKey" in err_text:
            deltas["sheet_rows_deduped"] += 1
            return ProcessOutcome(
                candidate=candidate,
                job_key=job_key,
                decision=decision,
                outcome="sheet_duplicate",
                job_url=job_url,
                error_reason=None,
                cover_letter_local_path=cover_letter_local_path,
                cover_letter_drive_file_id=cover_letter_drive_file_id,
                cover_letter_drive_folder_id=cover_letter_drive_folder_id,
                count_deltas=deltas,
                job_result={"job_key": job_key, "decision": decision, "status": "sheet_duplicate", "cover_letter": cover_link},
                errors=local_errors,
            )

        deltas["upload_errors"] += 1
        local_errors.append(f"sheet upload failed for `{job_key}`: {err_text or append_out.get('source')}")
        return ProcessOutcome(
            candidate=candidate,
            job_key=job_key,
            decision=decision,
            outcome="sheet_upload_error",
            job_url=job_url,
            error_reason=err_text or "sheet_upload_error",
            cover_letter_local_path=cover_letter_local_path,
            cover_letter_drive_file_id=cover_letter_drive_file_id,
            cover_letter_drive_folder_id=cover_letter_drive_folder_id,
            count_deltas=deltas,
            job_result={"job_key": job_key, "decision": decision, "status": "sheet_upload_error"},
            errors=local_errors,
        )

    if total_jobs_to_process > 0:
        available_slots: list[int] = list(range(1, process_workers + 1))
        outstanding: dict[Future[ProcessOutcome], tuple[SearchCandidate, int]] = {}
        next_idx = 0

        def _submit_until_full(executor: ThreadPoolExecutor) -> None:
            nonlocal next_idx
            while next_idx < total_jobs_to_process and len(outstanding) < process_workers and available_slots:
                slot = available_slots.pop(0)
                cand = discovered[next_idx]
                next_idx += 1
                _set_slot_state(
                    slot=slot,
                    state="running",
                    step_key="gather_info",
                    step_label="Gather info",
                    step_index=1,
                    candidate=cand,
                    job_key=cand.job_key,
                    emit_progress=True,
                )
                fut = executor.submit(_process_candidate, cand, slot=slot)
                outstanding[fut] = (cand, slot)

        try:
            with ThreadPoolExecutor(max_workers=process_workers, thread_name_prefix="indeed-process") as executor:
                _submit_until_full(executor)
                while outstanding:
                    done, _ = wait(set(outstanding.keys()), return_when=FIRST_COMPLETED)
                    for fut in done:
                        fallback_candidate, slot = outstanding.pop(fut)
                        try:
                            outcome = fut.result()
                        except Exception as exc:  # pragma: no cover - defensive runtime guard
                            err = f"worker exception for `{fallback_candidate.job_key}`: {exc}"
                            outcome = ProcessOutcome(
                                candidate=fallback_candidate,
                                job_key=fallback_candidate.job_key,
                                decision=DECISION_SKIP,
                                outcome="worker_exception",
                                job_url=_extract_job_url(fallback_candidate.job_listing),
                                error_reason="worker_exception",
                                cover_letter_local_path=None,
                                cover_letter_drive_file_id=None,
                                cover_letter_drive_folder_id=None,
                                count_deltas={
                                    "recommended_apply": 0,
                                    "recommended_maybe": 0,
                                    "skipped": 1,
                                    "extraction_errors": 0,
                                    "decision_errors": 1,
                                    "cover_letter_errors": 0,
                                    "upload_errors": 0,
                                    "sheet_rows_written": 0,
                                    "sheet_rows_deduped": 0,
                                },
                                job_result={"job_key": fallback_candidate.job_key, "decision": DECISION_SKIP, "status": "worker_exception"},
                                errors=[err],
                            )
                        for key, value in outcome.count_deltas.items():
                            if key in counts:
                                counts[key] += int(value)
                        errors.extend(outcome.errors)
                        job_results.append(outcome.job_result)
                        processed_jobs += 1
                        _set_slot_state(
                            slot=slot,
                            state="idle",
                            step_key="idle",
                            step_label="Idle",
                            step_index=0,
                            candidate=None,
                            job_key="",
                            emit_progress=False,
                        )
                        available_slots.append(slot)
                        available_slots.sort()
                        _emit_job_result(
                            candidate=outcome.candidate,
                            job_key=outcome.job_key,
                            decision=outcome.decision,
                            outcome=outcome.outcome,
                            job_url=outcome.job_url,
                            error_reason=outcome.error_reason,
                            cover_letter_local_path=outcome.cover_letter_local_path,
                            cover_letter_drive_file_id=outcome.cover_letter_drive_file_id,
                            cover_letter_drive_folder_id=outcome.cover_letter_drive_folder_id,
                        )
                    _submit_until_full(executor)
        finally:
            db_write_queue.join()
            sheet_write_queue.join()
            db_write_queue.put(db_stop_token)
            sheet_write_queue.put(sheet_stop_token)
            db_write_queue.join()
            sheet_write_queue.join()
            db_writer.join(timeout=2.0)
            sheet_writer.join(timeout=2.0)
    else:
        db_write_queue.put(db_stop_token)
        sheet_write_queue.put(sheet_stop_token)
        db_write_queue.join()
        sheet_write_queue.join()
        db_writer.join(timeout=2.0)
        sheet_writer.join(timeout=2.0)

    trigger = _coerce_text(payload.get("trigger")) or "scheduled"
    summary = (
        f"indeed_daily_search done trigger={trigger} "
        f"queries={counts['searched']} new_jobs={counts['new_jobs']} "
        f"apply={counts['recommended_apply']} maybe={counts['recommended_maybe']} "
        f"skip={counts['skipped']} rows_written={counts['sheet_rows_written']} "
        f"errors={len(errors)}"
    )
    snapshot = {
        "updated_at": _iso_now(),
        "summary": summary,
        "config": {
            "cover_letter_destination_path": destination_path,
            "cover_letter_destination_folder_id": destination_folder_id,
            "seen_ids_limit": seen_limit,
        },
        "counts": counts,
        "search_stats": {
            "queries_total": search_stats.get("queries_total"),
            "queries_ok": search_stats.get("queries_ok"),
            "jobs_returned_total": search_stats.get("jobs_returned_total"),
            "jobs_filtered_seen": search_stats.get("jobs_filtered_seen"),
            "jobs_new_total": search_stats.get("jobs_new_total"),
            "initial_seen_loaded": search_stats.get("initial_seen_loaded"),
        },
        "errors": errors[:20],
    }
    try:
        _upsert_task_state_snapshot(
            task_id=task_id,
            state_key="last_run_snapshot",
            value=snapshot,
            updated_by="indeed_daily_search_task",
            db_path=db_path,
        )
    except Exception as exc:  # pragma: no cover
        errors.append(f"failed to persist last_run_snapshot: {exc}")

    final_slot_seq, final_slots = _worker_slots_snapshot()
    progress.emit(
        {
            "stage": "done",
            "query_index": int(search_stats.get("queries_total") or 0),
            "query_total": int(search_stats.get("queries_total") or 0),
            "job_index": max(0, total_jobs_to_process),
            "job_total": max(0, total_jobs_to_process),
            "search_fraction": 1.0,
            "processed_jobs": processed_jobs,
            "total_jobs_for_processing": total_jobs_to_process,
            "worker_slots": final_slots,
            "slot_update_seq": final_slot_seq,
            "total_percent": 100.0,
            "status_line": summary,
        }
    )

    return {
        "ok": True,
        "summary": summary,
        "counts": counts,
        "search_stats": search_stats,
        "job_results": job_results,
        "errors": errors,
    }
